from docx import Document as DocxDocument
from typing import List
import io, os, tempfile
from langchain_core.document_loaders import BaseLoader
from langchain_core.documents import Document
import pymupdf
from docling.document_converter import DocumentConverter




class DOCXLoader(BaseLoader):
    """Loads a DOCX file (or stream) into a list of documents with metadata."""

    #Constructor for the DOCXLoader
    def __init__(
        self, 
        file_content: bytes, 
        source: str = "", 
        title: str = ""
    ): 
        """
        Initialize the loader with file content as bytes.
        :param file_content: Raw DOCX content as bytes (for API responses, etc.)
        """

        if file_content is None:
            raise ValueError("file_content must be provided.")
        
        self.file_stream = io.BytesIO(file_content) #This creates an in-memory binary stream from the bytes, which many file-processing libraries can work with as if it were a file on disk.
        self.docx_document = DocxDocument(self.file_stream)
        self.source = source
        self.title = title
        
            

    def load(self) -> List[Document]:
        """
        Load the DOCX and extract paragraphs as Document objects with metadata.
        :return: A list of Document objects.
        """
        documents = []
        metadata = {
            "source": self.source,  # Source will be set dynamically from caller
            'title': self.title,
            "document_type": "docx"
        }

        # Iterate over each paragraph in the DOCX
        for paragraph in self.docx_document.paragraphs:
            text = paragraph.text.strip()

            if text:  # Ignore empty paragraphs
                document = Document(
                    page_content=text, 
                    metadata=metadata
                )
                documents.append(document)

        return documents



class PDFLoader(BaseLoader):
    """ Loads a PDF file (or stream) into a list of LangChain Documents. """

    def __init__(
        self,
        file_content: bytes,
        source: str = "",
        title: str = ""
    ):

        if file_content is None:

            raise ValueError("file_content must be provided.")

        self.file_stream = io.BytesIO(file_content)       
        self.pdf_document = pymupdf.open(
            stream=self.file_stream.read(),
            filetype="pdf"
        )

        self.source = source
        self.title = title

    def _clean_text(self, text: str) -> str:

        lines = []

        for line in text.splitlines():
            line = line.strip()
            if line:
                lines.append(line)

        return "\n".join(lines)

    def _extract_page_text(self, page) -> str:

        text = page.get_text("text")

        return self._clean_text(text)

    def load(self) -> List[Document]:

        documents = []

        for page_number in range(len(self.pdf_document)):

            page = self.pdf_document[page_number]
            text = self._extract_page_text(page)

            # Future OCR hook
            if not text:
                # text = self._ocr_page(page)
                continue

            metadata = {
                "source": self.source,
                "title": self.title,
                "page": page_number + 1,
                "document_type": "pdf"
            }

            document = Document(
                page_content=text,
                metadata=metadata
            )

            documents.append(document)

        return documents



class DoclingLoader(BaseLoader):

    def __init__(
        self,
        file_content: bytes,
        filename: str
    ):

        self.file_content = file_content
        self.filename = filename

    def load(self):

        converter = DocumentConverter()

        with tempfile.NamedTemporaryFile(
            suffix=self.filename[
                self.filename.rfind("."):
            ],
            delete=False
        ) as temp:

            temp.write(self.file_content)
            temp_path = temp.name

        result = converter.convert(temp_path)

        markdown_text = (result.document.export_to_markdown())

        documents = [
            Document(
                page_content=markdown_text,
                metadata={
                    "source": self.filename
                        }
                    )
                ]
        
        if temp_path and os.path.exists(temp_path):
            os.remove(temp_path)

        return documents



def get_loader(filename: str,file_content: bytes):

    extension = filename.lower().split(".")[-1]

    if extension == "docx":

        return DOCXLoader(
            file_content=file_content
        )

    elif extension == "pdf":

        return DoclingLoader(
            file_content=file_content
        )

    raise ValueError(
        f"Unsupported file type: {extension}"
    )
import logging
from pydantic import BaseModel, FilePath
import sys
import os
from typing import Tuple, List, Optional, Union
from fastapi import UploadFile
from fastapi.datastructures import Headers
from pathlib import Path
import mimetypes
from zipfile import ZipFile


logger = logging.getLogger(__name__)
#########################################################################################################################
#pydantic models are used in FastAPI and other systems to: validate incoming data, ensure data types are correct, create structured request/response objects

class FileItem(BaseModel):
    path: str
    FileName: str
    is_public: int

class FileRequest(BaseModel):
    files_paths: List[FileItem]
    company_id: int
    files_title: str


class UserQuestion(BaseModel):
    user_text_question: str = ''
    #user_audio_question_path: str = ''
    stream: bool = False

#########################################################################################################################

from docx import Document as DocxDocument
from typing import List, Optional
import io
from langchain_core.document_loaders import BaseLoader
from langchain_core.documents import Document


class DOCXLoader(BaseLoader):
    """Loads a DOCX file (or stream) into a list of documents with metadata."""

    def __init__(self, file_content: Optional[bytes] = None): #constructor for the DOCXLoader
        """
        Initialize the loader with file content as bytes.
        :param file_content: Raw DOCX content as bytes (for API responses, etc.)
        """
        if file_content is not None:
            self.file_stream = io.BytesIO(file_content) #This creates an in-memory binary stream from the bytes, which many file-processing libraries can work with as if it were a file on disk.
            self.docx_document = DocxDocument(self.file_stream)
        else:
            raise ValueError("file_content must be provided.")

    def load(self) -> List[Document]:
        """
        Load the DOCX and extract paragraphs as Document objects with metadata.
        :return: A list of Document objects.
        """
        documents = []
        metadata = {
            "source": " ",  # Source will be set dynamically from caller
            'title': ' ',
            "page": 1
        }

        # Iterate over each paragraph in the DOCX
        for paragraph in self.docx_document.paragraphs:
            text = paragraph.text.strip()

            if text:  # Ignore empty paragraphs
                document = Document(page_content=text, metadata=metadata)
                documents.append(document)

        return documents

#########################################################################################################################
# -------------------------------------------------------------------------------------------------
#                           Load Documents
# -------------------------------------------------------------------------------------------------
async def load_documents_from_files_or_zip(
        uploaded_files: Union[bytes, list[UploadFile]],
        file_path: list[str],
        FileNames: list[str],
        #is_public: list[int],
        file_hashes: list[str]
        #file_paths_user: List[str]
) -> Tuple[list[Document], list[BaseLoader]]:
    """
    Load documents from a zip file or a list of uploaded files.

    Parameters:
    uploaded_files (Union[bytes, List[UploadFile]]): Either a single zip file content (as bytes) or a list of uploaded files.

    Returns:
    Tuple[List[Document], List[BaseLoader]]: A list of loaded documents and loaders used.
    """

    loaders = []
    all_documents = []
    # def process_file(file_content, file_path, filename, is_public, file_paths_user):
    def process_file(file_content, file_path, filename, file_hash):
        documents = []
        file_ext = Path(file_path).suffix.lower()
        print()

        
        if file_ext == ".docx":
            loader = DOCXLoader(file_content)
            documents.extend(loader.load())
            for doc in documents:
                doc.metadata["source"] = filename
                doc.metadata["hash"] = file_hash
                #doc.metadata["path"] = file_paths_user
            loaders.append(loader)
            all_documents.extend(documents)

    # If the uploaded file is a zip file (passed as bytes)
    if isinstance(uploaded_files, bytes):
        with ZipFile(io.BytesIO(uploaded_files)) as z:
            for file_name in z.namelist():
                print(file_name)
                file_ext = Path(file_name).suffix.lower()

                # Only process specific file types
                if file_ext in [".pdf", ".docx", ".csv", ".txt", ".xlsx", ".PNG"]:
                    with z.open(file_name) as file:
                        file_content = file.read()
                        # process_file(file_content, file_path[0], file_name, is_public[0], file_paths_user[0])
                        process_file(file_content, file_path[0], file_name, file_hashes[0])

    # If the uploaded files are individual files
    else:
        for i, file in enumerate(uploaded_files):
            file_content = await file.read()
            # print(file_content)
            # process_file(file_content, file_path[i], FileNames[i], is_public[i], file_paths_user[i])
            process_file(file_content, file_path[i], FileNames[i], file_hashes[i])

    return all_documents, loaders